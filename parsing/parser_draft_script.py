########################################################
# Xml files : extract, read, write
########################################################

# specific to extracting information from word documents
import os
import zipfile
#other tools useful in extracting the information from our document
import re
#to pretty print our xml:
import xml.dom.minidom
# to write to file in utf8
import io


def extract_xml_to_file(docx_path, xml_file, outfile_path):
    document = zipfile.ZipFile(docx_path)
    uglyXml = xml.dom.minidom.parseString(document.read(xml_file)).toprettyxml(indent='  ')
    with io.open(outfile_path, 'w', encoding='utf8') as f:
        f.write(uglyXml)

'''
extract_xml_to_file(
    '../../Downloads/questionnaires/CComptes_CCG_CH_SOM_v15032019.docx',
    'word/document.xml',
#    'word/styles.xml',
    "../../Downloads/questionnaires/CH_SOM.xml"
)
'''

########################################################
# Python-docx
########################################################

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

XML_TAG_ROOT = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


#document = Document('../../Downloads/questionnaires/0_1er questionnaire DDFiP Mayenne.docx')
#document = Document('../../Downloads/questionnaires/2190048_Quest.n°1 RCU_FontenaySBois.docx')
#document = Document('../../Downloads/questionnaires/Mayenne_formatted.docx')
#document = Document('../../Downloads/questionnaires/Fontenay.docx')
#document = Document('../../Downloads/questionnaires/Questionnaire_ADEME.1.docx')
#document = Document('../../Downloads/questionnaires/Questionnaire_N1_OPH_Bobigny.docx')
#document = Document('../../Downloads/questionnaires/CComptes_CCG_CH_SOM_v15032019.docx')
document = Document('../../Downloads/questionnaires/MTES.docx')


def print_obj(obj):
    for prop in dir(obj):
        if prop.startswith('__'): continue
        print(prop, end =" ")
    print()

def _find_size_in_params(params): # params can be rPr or pPr
    def find_val(sz):
        return float(sz.attrib[XML_TAG_ROOT + 'val'])

    sz = params.find(XML_TAG_ROOT + 'sz')
    if sz is not None:
        return find_val(sz) / 2
    sz = params.find(XML_TAG_ROOT + 'szCs')
    if sz is not None:
        return find_val(sz) / 2
    return None


def find_paragraph_size(paragraph, styles, verbose=False):
    def find_paragraph_params(paragraph):
        return paragraph._element.get_or_add_pPr()

    def find_run_params_in_paragraph_params(paragraph_params):
        for child in paragraph_params:
            if child.tag == XML_TAG_ROOT + 'rPr':
                return child

    pPr = find_paragraph_params(paragraph)
    rPr = find_run_params_in_paragraph_params(pPr)

    # try in the rPr contained in the pPr
    if rPr is None:
        if verbose: print('no pPr.rPr')
    else:
        size = _find_size_in_params(rPr)
        if size is not None:
            if verbose: print('found size in pPr.rPr')
            return size

    # try in the pPr
    size = _find_size_in_params(pPr)
    if size is not None:
        if verbose: print('found size in pPr')
        return size

    # try in the style defined for the paragraph -> last
    size = find_style_size(styles, paragraph.style.style_id, verbose)
    if size is not None:
        if verbose: print('found size in paragraph style')
        return size


def find_run_size(run):
    def find_run_params(run):
        return run.element.get_or_add_rPr()

    rPr = find_run_params(run)
    if rPr is None:
        return None
    return _find_size_in_params(rPr)

def find_style_size(styles, style_id, verbose=False):
    def find_style_by_id(styles, style_id):
        for s in styles:
            if s.style_id == style_id:
                return s

    def find_run_params(style):
        return style.element.get_or_add_pPr()

    style = find_style_by_id(styles, style_id)

    # try python-docx api (simple case only)
    if style.font.size is not None:
        if verbose: print('Found size in style.font')
        return style.font.size.pt

    # try in rPr
    rPr = find_run_params(style)
    if verbose: print('Found style.rPr ', rPr)
    sz = _find_size_in_params(rPr)
    if sz is not None:
        if verbose: print('Found size in style.rPr')
        return sz

    # try in base_style
    if style.base_style is not None:
        if verbose: print('Finding size in base_style : ', style.base_style)
        return find_style_size(styles, style.base_style.style_id, verbose)

    # try document defaults
    def find_default_rpr(styles):
        docDefaults = styles.element.find(XML_TAG_ROOT + 'docDefaults')
        rprDefaults = docDefaults.find(XML_TAG_ROOT + 'rPrDefault')
        return rprDefaults.find(XML_TAG_ROOT + 'rPr')

    default_rPr = find_default_rpr(styles)
    if default_rPr is not None:
        if verbose: print('Found default rPr')
        size = _find_size_in_params(default_rPr);
        if size is not None:
            if verbose: print('Found size in default rPr')
            return size

    if verbose: print('No style found')
    return None

def find_styles_in_document(document, verbose=False):
    styles = document.styles
    out = []
    for idx, p in enumerate(document.paragraphs):
        if verbose: print('--------', idx)
        if len(p.text) == 0 or p.text.isspace():
            continue
        if verbose: print(p.text)
        # try at the paragraph level
        size = find_paragraph_size(p, styles, verbose)
        if size is not None:
            if verbose: print('--', size, '---------------------HEADING' if size >= 14 else '')
            out.append({"text": p.text, "size": size})
            continue

        # try at the run level (first run only, should be enough)
        size = find_run_size(p.runs[0])
        if size is not None:
            if verbose: print('--', size, '---------------------HEADING' if size >= 14 else '')
            out.append({"text": p.text, "size": size})
            continue

        out.append({"text": p.text, "size": None})
        if verbose: print('----------------------NO SIZE FOUND')

    return out

styled = find_styles_in_document(document, verbose=True)


